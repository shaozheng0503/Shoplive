"""Generate the competition 作品信息摘要 document (Word)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = Path(__file__).parent / "ShopLive_作品信息摘要.docx"


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_run_cn(paragraph, text, font_name="宋体", size=Pt(10.5), bold=False):
    r = paragraph.add_run(text)
    r.font.size = size
    r.font.name = font_name
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    r.bold = bold
    return r


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_cn(p, "中国大学生计算机设计大赛", "黑体", Pt(22), bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_cn(p, "作品信息摘要", "黑体", Pt(20), bold=True)

    doc.add_paragraph()

    # ── Info table ──
    info_rows = [
        ("作品名称", "ShopLive — AI 驱动的电商短视频智能创作平台"),
        ("参赛小类", "Web 应用与开发"),
        ("作品编号", "（待填写）"),
        ("参赛学校", "（待填写）"),
        ("作者姓名", "（待填写）"),
        ("指导老师", "（待填写）"),
        ("联系电话", "（待填写）"),
        ("电子邮箱", "（待填写）"),
    ]
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_rows):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.width = Cm(3.5)
        c1.width = Cm(11.0)
        c0.text = ""
        c1.text = ""
        set_cell_shading(c0, "F2F6FC")
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_cn(p0, label, "黑体", Pt(11), bold=True)
        p1 = c1.paragraphs[0]
        add_run_cn(p1, value, "宋体", Pt(11))

    doc.add_paragraph()

    # ── 作品简介 ──
    p = doc.add_paragraph()
    add_run_cn(p, "一、作品简介", "黑体", Pt(14), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    add_run_cn(p, (
        "ShopLive 是一款面向中小电商商家的 AI 驱动短视频智能创作平台。"
        "平台解决了电商短视频制作成本高、周期长、创意门槛高三大核心痛点，"
        "将传统的「天级」视频制作周期压缩到「分钟级」。"
    ))

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    add_run_cn(p, (
        "用户只需提供商品链接或商品图片，系统即可自动完成商品信息解析、AI 脚本生成、"
        "多引擎视频合成的全流程。平台创新性地提供了「爆款视频复刻」功能——"
        "粘贴一条抖音/小红书爆款视频链接，系统自动进行 ASR 字幕转写、"
        "LLM 结构化节奏拆解、商品替换与多引擎提示词生成，帮助商家一键复用已验证的爆款节奏。"
    ))

    # ── 核心功能 ──
    p = doc.add_paragraph()
    add_run_cn(p, "二、核心功能与技术亮点", "黑体", Pt(14), bold=True)

    features = [
        "多引擎 AI 视频生成：集成 Veo 3.1、即梦 3.0、LTX 2.3、Grok 四种视频生成引擎，支持 4-24 秒视频",
        "爆款视频一键复刻：ASR 字幕提取 → LLM 结构化分析 → 4 引擎专属提示词 → 置信度评分 → 一键生成",
        "商品链接智能解析：输入商品 URL 自动爬取名称、价格、卖点、主图等结构化信息",
        "AI Agent 智能对话：支持 10 轮 Tool-Calling 自动工具调用，实现「对话即编辑」",
        "FFmpeg 在线视频编辑：加速/裁剪/文字叠加/多源合成，异步任务队列",
        "AI 商品图片生成：Gemini 驱动的图片生成与质量检测",
    ]
    for feat in features:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(20)
        add_run_cn(p, f"  - {feat}")

    # ── AI 大模型 ──
    p = doc.add_paragraph()
    add_run_cn(p, "三、AI 大模型应用", "黑体", Pt(14), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    add_run_cn(p, (
        "本项目深度集成多个 AI 大模型：(1) Google Gemini 2.5 Flash——用于视频多模态 ASR 字幕提取"
        "和商品图片分析，直接上传 mp4 视频进行理解，无需独立语音识别服务；"
        "(2) Vertex AI LLM——用于爆款视频结构化分析（输出严格 JSON）、商品脚本生成和提示词增强；"
        "(3) Google Veo 3.1——核心视频生成引擎，支持文本/图片到视频；"
        "(4) 即梦 3.0 / LTX 2.3——备选视频生成引擎，提供风格差异化选择。"
        "AI 大模型的应用贯穿商品解析、脚本生成、视频合成全链路，并通过 Prompt 工程、"
        "结构化输出约束、智能兜底机制和缓存策略实现工程化落地。"
    ))

    # ── 技术架构 ──
    p = doc.add_paragraph()
    add_run_cn(p, "四、技术架构", "黑体", Pt(14), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)
    add_run_cn(p, (
        "系统采用前后端分离的 B/S 架构。前端使用原生 JavaScript + HTML/CSS 实现轻量化 SPA，"
        "后端基于 Python Flask 框架提供 60 个 RESTful API 端点，"
        "AI 服务层对接 Google Vertex AI 生态及第三方引擎。"
        "项目代码量：后端 Python 约 16,700 行，前端 JS 约 14,100 行，"
        "自动化测试 370 个用例全部通过。"
    ))

    # ── 创新点 ──
    p = doc.add_paragraph()
    add_run_cn(p, "五、主要创新点", "黑体", Pt(14), bold=True)

    innovations = [
        "全自动爆款视频复刻管线（分享链接解析 → ASR → LLM 结构化分析 → 多引擎提示词），填补电商+AI 视频交叉领域空白",
        "多引擎提示词自适应：同一复刻方案自动为 Veo/即梦/LTX/Grok 生成风格差异化的专属提示词",
        "置信度评分机制：综合分析来源、字幕数量和分镜完整度，量化复刻方案可靠程度",
        "Agent 10 轮 Tool-Calling 循环 + SSE 实时推送，实现「对话即编辑」的视频创作范式",
        "字幕噪声过滤：自动识别并清除 ASR 中的 [music]/[背景音乐] 等非语音内容",
    ]
    for inno in innovations:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(20)
        add_run_cn(p, f"  - {inno}")

    # ── 关键词 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_cn(p, "关键词：", "黑体", Pt(11), bold=True)
    add_run_cn(p, "AI 视频生成；电商短视频；爆款视频复刻；大模型应用；Gemini；Veo；Agent Tool-Calling")

    doc.save(str(OUT))
    print(f"Document saved to: {OUT}")
    print(f"File size: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
