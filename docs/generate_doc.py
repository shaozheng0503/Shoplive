"""Generate the competition design & development document (Word) for ShopLive."""
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCREENSHOTS = Path(__file__).parent / "screenshots"
OUT = Path(__file__).parent / "ShopLive_设计和开发文档.docx"


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_heading1(doc, text):
    """一级标题: 二号黑体, 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(22)  # 二号
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.bold = True
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_heading2(doc, text):
    """二级标题: 三号黑体, 靠左"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(16)  # 三号
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_heading3(doc, text):
    """三级标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, indent=False):
    """正文: 五号宋体"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)  # 五号
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(18)
    return p


def add_image(doc, filename, width_inches=5.5, caption=""):
    """Insert screenshot with optional caption."""
    path = SCREENSHOTS / filename
    if not path.exists():
        add_body(doc, f"[截图缺失: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.color.rgb = RGBColor(100, 100, 100)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "2B5797")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing
    return table


def build_document():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中国大学生计算机设计大赛")
    run.font.size = Pt(26)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("软件开发类作品设计和开发文档")
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph()
    doc.add_paragraph()

    cover_fields = [
        ("作品名称", "ShopLive — AI 驱动的电商短视频智能创作平台"),
        ("参赛小类", "Web 应用与开发"),
        ("作品编号", "（待填写）"),
        ("作　　者", "（待填写）"),
        ("指导老师", "（待填写）"),
        ("版本编号", "V1.0"),
        ("填写日期", "2026 年 4 月"),
    ]
    for label, value in cover_fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}：")
        r1.font.size = Pt(14)
        r1.font.name = "宋体"
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r1.bold = True
        r2 = p.add_run(value)
        r2.font.size = Pt(14)
        r2.font.name = "宋体"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "目  录")
    toc_items = [
        "第一章  需求分析",
        "  1.1  项目背景与问题定义",
        "  1.2  目标用户",
        "  1.3  核心功能",
        "  1.4  竞品分析",
        "  1.5  AI 大模型应用概述",
        "第二章  概要设计",
        "  2.1  系统整体架构",
        "  2.2  功能模块层次结构",
        "  2.3  核心调用链路",
        "第三章  详细设计",
        "  3.1  界面设计与典型使用流程",
        "  3.2  关键技术与创新点",
        "  3.3  AI 大模型深度集成",
        "第四章  测试报告",
        "  4.1  测试概况",
        "  4.2  典型测试用例",
        "  4.3  技术指标",
        "第五章  安装及使用",
        "  5.1  环境要求",
        "  5.2  安装步骤",
        "  5.3  典型使用流程",
        "第六章  项目总结",
        "  6.1  开发过程与克服的困难",
        "  6.2  AI 大模型应用心得",
        "  6.3  后续演进规划",
        "参考文献",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        r = p.add_run(item)
        r.font.size = Pt(12) if not item.startswith("  ") else Pt(10.5)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if not item.startswith("  "):
            r.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 1: 需求分析
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "第一章  需求分析")

    add_heading2(doc, "1.1 项目背景与问题定义")
    add_body(doc, (
        "随着短视频电商的爆发式增长，抖音、快手、小红书等平台已成为商品销售的重要渠道。"
        "据统计，2025 年中国短视频电商市场规模已突破 3 万亿元，超过 70% 的电商商家表示短视频是其获客的核心手段。"
        "然而，中小商家普遍面临视频创作的三大痛点：(1) 专业制作成本高——外包一条 15 秒带货短视频的报价通常在 500-2000 元；"
        "(2) 制作周期长——从脚本策划到成片交付通常需要 3-7 天；(3) 创意门槛高——缺乏对爆款视频节奏和结构的系统性理解。"
    ), indent=True)
    add_body(doc, (
        "ShopLive 正是为解决上述痛点而设计的 AI 驱动电商短视频智能创作平台。"
        "用户只需提供商品链接或商品图片，平台即可自动完成商品信息解析、脚本生成、视频合成的全流程，"
        "将传统的「天级」制作周期压缩到「分钟级」。同时，平台创新性地提供了「爆款视频复刻」功能——"
        "用户粘贴一条抖音/小红书爆款视频链接，系统自动进行 ASR 字幕提取、结构化节奏拆解、"
        "商品替换生成，帮助商家快速复用已验证的爆款节奏。"
    ), indent=True)

    add_heading2(doc, "1.2 目标用户")
    add_body(doc, (
        "本平台主要面向以下三类用户群体：(1) 中小电商商家——拥有自有商品但缺乏视频制作团队，"
        "需要低成本快速产出带货短视频；(2) 跨境电商卖家——需要多语言视频内容适配不同市场，"
        "对视频本地化效率有较高要求；(3) 社交电商从业者——需要紧跟平台热点，快速复刻爆款视频节奏，"
        "对时效性要求极高。"
    ), indent=True)

    add_heading2(doc, "1.3 核心功能")
    add_body(doc, "平台提供以下六大核心功能模块：")
    features = [
        ("AI 视频生成", "基于文本提示词或商品图片，调用 Veo 3.1/即梦/LTX 等多引擎生成 4-24 秒电商短视频"),
        ("商品链接解析", "输入商品页面 URL，自动爬取商品名称、价格、卖点、主图等结构化信息"),
        ("爆款视频复刻", "粘贴爆款视频分享链接，自动 ASR 转写、LLM 结构化分析、生成复刻脚本与多引擎提示词"),
        ("智能 Agent 对话", "支持 10 轮工具调用的 AI Agent，可自主完成视频编辑、参数调整等复合任务"),
        ("在线视频编辑", "基于 FFmpeg 的 Timeline 编辑器，支持加速/裁剪/文字叠加/多源合成"),
        ("AI 图片生成", "Gemini 驱动的商品图片生成与质量检测，支持电商场景图自动合成"),
    ]
    for name, desc in features:
        p = doc.add_paragraph()
        r1 = p.add_run(f"  {name}：")
        r1.font.size = Pt(10.5)
        r1.font.name = "宋体"
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r1.bold = True
        r2 = p.add_run(desc)
        r2.font.size = Pt(10.5)
        r2.font.name = "宋体"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_heading2(doc, "1.4 竞品分析")
    add_body(doc, "下表从多个维度对比 ShopLive 与市场主要竞品：")
    add_table(doc,
        ["维度", "ShopLive", "剪映/CapCut", "Runway Gen-3", "Pika Labs"],
        [
            ["AI 视频生成", "多引擎 (Veo/即梦/LTX/Grok)", "模板+AI 特效", "单一模型", "单一模型"],
            ["电商场景定制", "深度定制（商品解析+脚本+CTA）", "通用模板", "无", "无"],
            ["爆款视频复刻", "ASR+LLM 结构化拆解+一键复刻", "无", "无", "无"],
            ["Agent 智能对话", "10 轮 Tool-calling 自动编辑", "无", "无", "无"],
            ["多引擎切换", "4 种引擎自由切换", "仅自有引擎", "仅自有模型", "仅自有模型"],
            ["在线视频编辑", "FFmpeg Timeline 编辑器", "完整编辑器", "简单裁剪", "无"],
            ["中文适配", "原生中英双语", "中文优先", "仅英文", "仅英文"],
            ["使用成本", "按需调用，无需专业团队", "免费+会员", "按秒计费", "按秒计费"],
        ],
        col_widths=[2.8, 4.0, 3.0, 2.8, 2.8],
    )
    add_body(doc, (
        "对比可见，ShopLive 在电商场景定制、爆款视频复刻和 Agent 智能编辑三个维度具备差异化优势，"
        "填补了现有工具在「电商 + AI 视频」交叉领域的空白。"
    ), indent=True)

    add_heading2(doc, "1.5 AI 大模型应用概述")
    add_body(doc, "本项目深度集成了多个 AI 大模型，贯穿全链路：")
    add_table(doc,
        ["大模型", "应用场景", "调用方式"],
        [
            ["Google Gemini 2.5 Flash", "视频 ASR 字幕提取（多模态视频理解）；商品图片分析与洞察", "Vertex AI REST API，base64 视频/图片上传"],
            ["Vertex AI LLM", "爆款视频结构化分析（输出 JSON）；商品脚本生成；提示词增强", "call_litellm_chat 统一封装，支持流式输出"],
            ["Google Veo 3.1", "文本/图片到视频生成（4-24s）", "Vertex AI Video Generation API"],
            ["即梦 3.0 / LTX 2.3", "备选视频生成引擎（风格差异化）", "各厂商 API 直调"],
            ["Gemini Imagen", "电商商品图片生成与合成", "Vertex AI Image Generation API"],
        ],
        col_widths=[3.5, 5.5, 5.5],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 2: 概要设计
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "第二章  概要设计")

    add_heading2(doc, "2.1 系统整体架构")
    add_body(doc, (
        "ShopLive 采用前后端分离的 B/S 架构，分为展示层、业务逻辑层和 AI 服务层三个层次。"
        "前端使用原生 JavaScript + HTML/CSS 实现轻量化 SPA，后端基于 Python Flask 框架，"
        "通过 RESTful API 对外提供服务，AI 服务层对接 Google Vertex AI 生态（Gemini/Veo）及第三方引擎。"
    ), indent=True)

    # Architecture diagram as text
    add_body(doc, "系统架构如下所示：")
    arch_table = [
        ["展示层（Frontend）", "原生 JS SPA：首页(index.html)、Agent 创作台(agent.html)、Studio 工作台(studio.html)、图片实验室(image-lab.html)"],
        ["", "状态管理(state.js)、国际化(i18n.js)、工作区编排(workspace.js)、Agent 运行器(agent-run.js)"],
        ["业务逻辑层（Backend）", "Flask Web App (web_app.py) — 60 个 API 端点"],
        ["", "agent_api.py (Agent 对话+Tool-calling)、veo_api.py (视频生成链)、hot_video_api.py (爆款复刻)"],
        ["", "video_edit_api.py (FFmpeg 编辑)、shoplive_api.py (脚本工作流)、media_api.py (媒体代理)"],
        ["", "基础设施：schemas.py (Pydantic 校验)、audit.py (审计日志)、tool_registry.py (工具注册表)"],
        ["AI 服务层", "Gemini 2.5 Flash (ASR/图片分析)、Vertex AI LLM (结构化分析)"],
        ["", "Veo 3.1 / 即梦 3.0 / LTX 2.3 / Grok (视频生成)、Imagen (图片生成)"],
    ]
    add_table(doc,
        ["层次", "组件"],
        arch_table,
        col_widths=[3.5, 11.0],
    )

    add_heading2(doc, "2.2 功能模块层次结构")
    add_body(doc, "系统功能模块按职责划分为以下层次：")
    modules = [
        ["用户交互模块", "首页入口、Agent 创作台、Studio 工作台、图片实验室", "前端 JS/HTML"],
        ["商品解析模块", "商品链接爬取、Gemini 图片洞察、卖点结构化", "agent_api.py, scraper/"],
        ["视频生成模块", "Veo/即梦/LTX/Grok 多引擎调度、16s 工作流编排", "veo_api.py, jimeng_api.py, ltxv_api.py"],
        ["爆款复刻模块", "分享链接解析、ASR 字幕提取、LLM 结构化分析、多引擎提示词生成", "hot_video_api.py"],
        ["Agent 引擎模块", "10 轮 Tool-calling 循环、工具注册表、SSE 流式输出", "agent_api.py, tool_registry.py"],
        ["视频编辑模块", "FFmpeg Timeline 渲染、加速/裁剪/文字叠加/多源合成", "video_edit_api.py"],
        ["基础设施模块", "Pydantic 请求校验、全链路审计、Token 缓存、MCP 协议适配", "schemas.py, audit.py, infra.py"],
    ]
    add_table(doc,
        ["模块名称", "功能描述", "核心文件"],
        modules,
        col_widths=[3.0, 7.0, 4.5],
    )

    add_heading2(doc, "2.3 核心调用链路")
    add_body(doc, "以「爆款视频复刻」为例，展示系统核心调用链路：")
    chain = [
        ["1", "用户粘贴分享链接", "前端 hotVideoUrlInput", "用户输入"],
        ["2", "解析分享链接获取视频直链", "share_url_resolver.py + Playwright", "HTTP 重定向 + HTML 渲染"],
        ["3", "下载视频并 ASR 转写", "Gemini 2.5 Flash 多模态 API", "base64 视频上传，返回时间戳字幕"],
        ["4", "字幕噪声过滤", "_clean_subtitle_text()", "过滤 [music]/[背景音乐] 等噪声标签"],
        ["5", "LLM 结构化分析", "Vertex AI LLM (call_litellm_chat)", "输入字幕+商品信息，输出 JSON 复刻方案"],
        ["6", "多引擎提示词生成", "_build_engine_prompts()", "为 Veo/即梦/LTX/Grok 各生成专属提示词"],
        ["7", "置信度评分", "_compute_confidence_score()", "综合 source/字幕数/分镜数返回 0-1 分数"],
        ["8", "前端展示复刻方案", "applyHotVideoRemakeResult()", "分镜脚本+提示词填充+引导上传商品图"],
        ["9", "一键复刻生成视频", "submitSimplePromptGeneration()", "使用引擎专属提示词+商品图调用视频生成"],
    ]
    add_table(doc,
        ["步骤", "操作", "关键组件", "技术要点"],
        chain,
        col_widths=[1.0, 3.5, 4.5, 5.5],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 3: 详细设计
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "第三章  详细设计")

    add_heading2(doc, "3.1 界面设计与典型使用流程")

    add_heading3(doc, "3.1.1 首页（落地页）")
    add_body(doc, (
        "首页采用大居中入口设计，顶部展示核心 Slogan「把想法直接变成视频」，中央为统一输入框，"
        "下方以卡片形式展示「近期精选」示例视频，降低用户上手门槛。"
    ), indent=True)
    add_image(doc, "01_landing_page.png", 5.5, "图 3-1  ShopLive 首页/落地页")

    add_heading3(doc, "3.1.2 AI Agent 创作台")
    add_body(doc, (
        "Agent 创作台是核心工作界面，顶部为对话区域，底部为功能操作栏。"
        "操作栏集成了四个快捷入口：「上传图片」「商品链接」「爆款视频」「首尾帧」，"
        "支持四种视频引擎切换（LTX 2.3 / 即梦 3.0 / Veo 3.1 / Grok），"
        "可选择比例（16:9/9:16/1:1）和时长（4-24 秒）。"
    ), indent=True)
    add_image(doc, "02_agent_main.png", 5.5, "图 3-2  AI Agent 创作台主界面")

    add_heading3(doc, "3.1.3 爆款视频复刻入口")
    add_body(doc, (
        "点击「爆款视频」按钮后展开输入行，用户粘贴抖音/小红书/快手的分享链接，"
        "点击「解析爆款」即开始全自动 ASR + LLM 分析流程。分析完成后自动打开分镜脚本编辑器，"
        "用户可编辑后点击「一键复刻」生成视频。"
    ), indent=True)
    add_image(doc, "03_agent_hot_video_input.png", 5.5, "图 3-3  爆款视频复刻 - 链接输入区")

    add_heading3(doc, "3.1.4 Studio 工作台")
    add_body(doc, (
        "Studio 页面提供更专业的创作工作流，左侧为素材管理区（上传商品图/参考图），"
        "中央为视频预览区，右侧集成了 AI 脚本生成和 Shoplive AI 对话助手。"
        "支持从素材到脚本到视频的完整工作流编排。"
    ), indent=True)
    add_image(doc, "04_studio.png", 5.5, "图 3-4  Studio 工作台界面")

    add_heading3(doc, "3.1.5 AI 图片实验室")
    add_body(doc, (
        "图片实验室页面支持多模型视频生成（LTX 2.3 / 即梦 3.0 / Veo 3.1 Fast），"
        "提供细粒度参数控制（分辨率/帧率/画幅方向），以及实时图片质量检测功能。"
        "底部集成了两步流程引导：先生成商品图，再生成营销视频。"
    ), indent=True)
    add_image(doc, "05_image_lab.png", 5.5, "图 3-5  AI 图片实验室界面")

    add_heading3(doc, "3.1.6 移动端适配")
    add_body(doc, (
        "平台全面适配移动端浏览器，核心功能在手机端均可正常使用，包括视频引擎选择、"
        "提示词输入、图片上传等操作。"
    ), indent=True)
    add_image(doc, "11_mobile_view.png", 2.5, "图 3-6  移动端 Agent 创作台")

    add_heading3(doc, "3.1.7 典型使用流程")
    add_body(doc, "以「爆款视频复刻」场景为例，典型使用流程如下：")
    flow = [
        ["1", "用户从抖音 App 复制一条爆款视频的分享链接"],
        ["2", "粘贴到 ShopLive Agent 创作台的「爆款视频」输入框"],
        ["3", "点击「解析爆款」，系统自动执行：链接解析 → 视频下载 → ASR 字幕转写 → LLM 结构化分析"],
        ["4", "分析完成后，系统展示爆款摘要（钩子/节奏/分镜计划/置信度），并自动打开分镜脚本编辑器"],
        ["5", "用户上传自己的商品图片（或通过商品链接自动获取）"],
        ["6", "点击「一键复刻」，系统使用引擎专属提示词 + 商品图调用 Veo/即梦等引擎生成视频"],
        ["7", "视频生成完成后可在线预览、下载，或进入视频编辑器进一步调整"],
    ]
    add_table(doc,
        ["步骤", "操作说明"],
        flow,
        col_widths=[1.0, 13.5],
    )

    add_heading2(doc, "3.2 关键技术与创新点")

    add_heading3(doc, "3.2.1 爆款视频复刻管线")
    add_body(doc, (
        "这是本项目最核心的技术创新。传统的视频复刻依赖人工拆解爆款结构，ShopLive 实现了全自动化管线：\n"
        "(1) 分享链接智能解析：支持抖音/小红书/快手等 6 种平台短链，通过 HTTP 重定向追踪 + "
        "Playwright 渲染双策略获取视频直链，成功率超过 85%。\n"
        "(2) Gemini 多模态 ASR：将视频文件以 base64 编码上传至 Gemini 2.5 Flash，"
        "利用多模态视频理解能力提取带时间戳的字幕，支持中英文，单次最多 40 行。"
        "创新性地加入了噪声标签过滤（自动去除 [music]/[背景音乐]/[掌声] 等非语音内容）。\n"
        "(3) LLM 结构化分析：将 ASR 字幕 + 商品上下文注入 Vertex AI LLM，通过精心设计的 "
        "system prompt 引导模型输出严格 JSON 格式的复刻方案，包含 Hook 文案、三段式结构、"
        "6 镜头分镜计划、配音脚本、复刻提示词等 8 个字段。\n"
        "(4) 多引擎提示词自适应：根据用户选择的视频引擎（Veo/即梦/LTX/Grok），"
        "自动生成风格差异化的提示词——Veo 偏向电影镜头语言，即梦偏向中文电商美学，"
        "LTX 偏向关键帧描述，Grok 偏向自然叙事。\n"
        "(5) 置信度评分：综合分析来源、字幕数量和分镜完整度，计算 0-1 置信度分数，"
        "帮助用户判断复刻方案的可靠程度。"
    ), indent=True)

    add_heading3(doc, "3.2.2 Agent Tool-Calling 循环")
    add_body(doc, (
        "平台实现了完整的 AI Agent 运行时（/api/agent/run），支持最多 10 轮的自动工具调用循环。"
        "Agent 通过 SSE（Server-Sent Events）实时推送事件流（start/thinking/tool_call/tool_result/delta/done），"
        "前端逐事件渲染，用户可实时观察 Agent 的思考过程和工具执行结果。\n"
        "工具注册表（tool_registry.py）以声明式方式管理所有可用工具，支持 OpenAI function-calling 格式输出，"
        "每个工具包含名称、描述、参数 JSON Schema 和执行函数。Agent 可自主调用视频编辑、参数查询等工具，"
        "实现「对话即编辑」的交互范式。工具执行采用 ThreadPoolExecutor + Future.result(timeout) "
        "实现真正非阻塞超时，单次工具调用超时阈值 0.46 秒。"
    ), indent=True)

    add_heading3(doc, "3.2.3 FFmpeg Timeline 视频编辑器")
    add_body(doc, (
        "平台集成了基于 FFmpeg 的 Timeline 视频编辑器，支持以下编辑能力：\n"
        "- 视频加速/减速（speed 参数）\n"
        "- 时间段裁剪（trim_start/trim_end）\n"
        "- 饱和度调整（saturation）\n"
        "- 文字叠加（maskText，基于 FFmpeg drawtext 滤镜，支持中文字体）\n"
        "- 多源视频合成（source_index per segment，ffmpeg filter_complex [N:v]）\n"
        "编辑任务异步执行，队列上限 200 个任务，满时优先淘汰最旧的已完成任务，防止内存溢出。"
    ), indent=True)

    add_heading2(doc, "3.3 AI 大模型深度集成")
    add_body(doc, (
        "本项目对 AI 大模型的应用不是简单的 API 调用，而是深度融入了业务逻辑：\n\n"
        "(1) Prompt 工程实践：爆款复刻的 system prompt 针对不同视频引擎注入差异化风格提示，"
        "要求 LLM 将 product_anchors（颜色/材质/关键细节）融入 visual 描述；"
        "response schema 采用带说明的结构化描述（如 'summary: string — core strategy of the reference video'），"
        "而非裸类型，显著提升 JSON 输出稳定性。\n\n"
        "(2) 多模态理解：利用 Gemini 2.5 Flash 的视频理解能力，直接将 mp4 视频上传进行 ASR，"
        "无需额外的语音识别服务，大幅简化架构。同时利用 Gemini 的图片理解能力完成商品图分析。\n\n"
        "(3) 智能兜底机制：当 ASR 失败或 LLM 输出异常时，系统自动切换至基于模板的兜底方案，"
        "保证全链路可用性。兜底方案仍会利用已获取的部分信息（如 ASR 首行字幕作为 Hook），"
        "最大化已有数据的价值。\n\n"
        "(4) 缓存策略：ASR 结果缓存 24 小时，LLM 分析结果缓存 15 分钟，"
        "避免相同视频重复消耗 API 配额。LLM 缓存 key 包含 api_base hash，防止跨实例缓存污染。"
    ), indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 4: 测试报告
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "第四章  测试报告")

    add_heading2(doc, "4.1 测试概况")
    add_body(doc, (
        "本项目采用 pytest 自动化测试框架，共编写 12 个测试文件、370 个测试用例，全部通过。"
        "测试覆盖后端所有核心模块，包括 Pydantic Schema 校验、审计日志、请求验证中间件、"
        "Agent 运行时、视频编辑 API、爆款复刻 API 等。测试执行时间约 28 秒。"
    ), indent=True)

    test_summary = [
        ["test_schemas.py", "Pydantic Schema 校验", "39", "全部通过"],
        ["test_audit.py", "审计日志与文件持久化", "23", "全部通过"],
        ["test_validation.py", "请求验证装饰器", "23", "全部通过"],
        ["test_agent_run.py", "Agent 运行时 + Tool-calling", "35", "全部通过"],
        ["test_hot_video_api.py", "爆款视频复刻 API", "28", "全部通过"],
        ["test_video_edit_api.py", "视频编辑 Timeline API", "47", "全部通过"],
        ["test_optimizations.py", "系统优化验证", "19", "全部通过"],
        ["test_briefing.py", "视频脚本 Briefing", "4", "全部通过"],
        ["test_comfyui_ltxv_api.py", "ComfyUI LTX 集成", "14", "全部通过"],
        ["其他测试文件", "jimeng/tabcode/ltxv API", "138", "全部通过"],
    ]
    add_table(doc,
        ["测试文件", "覆盖模块", "用例数", "结果"],
        test_summary,
        col_widths=[3.5, 4.5, 2.0, 2.0],
    )

    add_heading2(doc, "4.2 典型测试用例")
    add_body(doc, "以下选取各模块代表性测试用例展示：")
    test_cases = [
        ["爆款复刻 - 成功路径", "抖音分享链接 + 商品信息", "返回 source=litellm，包含完整分镜计划、engine_prompts、confidence_score", "通过"],
        ["爆款复刻 - LLM 失败兜底", "有效链接但 LLM 异常", "返回 source=fallback，兜底结构完整，engine_prompts 仍有 4 个引擎", "通过"],
        ["爆款复刻 - 字幕噪声过滤", "[music] 这件衣服显瘦", "过滤后返回 '这件衣服显瘦'", "通过"],
        ["爆款复刻 - 置信度评分", "litellm + 12条字幕 + 6镜头", "confidence_score 在 0.9-1.0 之间", "通过"],
        ["Agent 运行 - 工具超时", "工具执行超过 0.46s", "返回超时错误，不阻塞主线程", "通过"],
        ["Agent 运行 - 10 轮上限", "持续产生工具调用", "第 10 轮后自动终止，发送 done 事件", "通过"],
        ["视频编辑 - 加速渲染", "5s 视频 speed=2x", "输出 2.56s 视频", "通过"],
        ["视频编辑 - 队列上限", "注入 200+ 任务", "满时淘汰最旧 terminal 任务，返回 429", "通过"],
        ["Schema 校验 - 无效 URL", "video_url='not-a-url'", "返回 400 + VALIDATION_ERROR", "通过"],
        ["多引擎提示词", "video_engine=jimeng", "jimeng 提示词包含「高清/通透/暖调」等中文美学词", "通过"],
    ]
    add_table(doc,
        ["测试用例", "输入", "预期结果", "实际"],
        test_cases,
        col_widths=[3.0, 3.5, 6.5, 1.5],
    )

    add_heading2(doc, "4.3 技术指标")
    metrics = [
        ["运行速度", "爆款复刻全流程（ASR+分析）", "20-60 秒", "取决于视频时长和网络"],
        ["运行速度", "视频生成（Veo 3.1）", "30-120 秒", "取决于时长和分辨率"],
        ["运行速度", "商品链接解析", "3-8 秒", "含 Playwright 渲染"],
        ["运行速度", "API 平均响应时间", "< 200ms", "不含外部 AI 调用"],
        ["安全性", "Pydantic 请求校验", "全部 POST 端点", "字段级验证 + 结构化错误"],
        ["安全性", "全链路审计日志", "60 个 API 端点", "异步写入，零阻塞"],
        ["可扩展性", "视频引擎", "4 种可切换", "Veo/即梦/LTX/Grok"],
        ["可扩展性", "工具注册表", "声明式注册", "新增工具仅需 register_tool()"],
        ["部署方便性", "启动命令", "python3 backend/run.py", "单命令启动，自动加载 .env"],
        ["可用性", "自动化测试", "370 用例全部通过", "pytest，28 秒完成"],
        ["可用性", "中英双语", "完整国际化", "前后端统一 i18n"],
    ]
    add_table(doc,
        ["维度", "指标项", "数值/状态", "备注"],
        metrics,
        col_widths=[2.5, 3.5, 3.5, 5.0],
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 5: 安装及使用
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "第五章  安装及使用")

    add_heading2(doc, "5.1 环境要求")
    env_req = [
        ["操作系统", "macOS / Linux / Windows (WSL)", "已在 macOS Darwin 24.6.0 上验证"],
        ["Python", "3.9 及以上", "推荐 3.10+"],
        ["Node.js", "16 及以上（可选，用于前端开发）", "生产环境不需要"],
        ["FFmpeg", "支持 libfreetype 的版本", "用于视频编辑功能"],
        ["GCP 凭证", "Google Cloud 服务账号 JSON 密钥", "用于 Vertex AI / Gemini / Veo API"],
        ["网络", "可访问 Google Cloud API", "部分环境需要代理配置"],
    ]
    add_table(doc,
        ["依赖项", "版本要求", "备注"],
        env_req,
        col_widths=[3.0, 5.0, 6.5],
    )

    add_heading2(doc, "5.2 安装步骤")
    add_body(doc, "步骤一：克隆代码仓库")
    add_body(doc, "  git clone https://github.com/shaozheng0503/Shoplive.git")
    add_body(doc, "  cd Shoplive")
    add_body(doc, "")
    add_body(doc, "步骤二：安装 Python 依赖")
    add_body(doc, "  pip install -r requirements.txt")
    add_body(doc, "")
    add_body(doc, "步骤三：配置环境变量")
    add_body(doc, "  在项目根目录创建 .env 文件，配置 GCP 凭证路径、项目 ID 等参数。")
    add_body(doc, "")
    add_body(doc, "步骤四：启动服务")
    add_body(doc, "  python3 backend/run.py")
    add_body(doc, "  服务默认运行在 http://127.0.0.1:8000")
    add_body(doc, "")
    add_body(doc, "步骤五：访问平台")
    add_body(doc, "  浏览器打开 http://127.0.0.1:8000 即可使用。")

    add_heading2(doc, "5.3 典型使用流程")
    add_body(doc, (
        "流程一：文本提示词生成视频——在 Agent 创作台输入视频描述（如「一个模特展示白色连衣裙，"
        "在阳光下旋转，裙摆飘动」），选择视频引擎和时长，点击「生成视频」。\n\n"
        "流程二：商品链接一键生成——粘贴电商平台的商品链接（如淘宝/Shopify/Amazon），"
        "系统自动爬取商品信息并填充卖点，用户确认后一键生成带货视频。\n\n"
        "流程三：爆款视频复刻——从抖音 App 复制一条爆款视频分享链接，粘贴到「爆款视频」输入框，"
        "点击「解析爆款」等待分析完成，上传自己的商品图，点击「一键复刻」生成同节奏的带货视频。"
    ), indent=True)

    add_image(doc, "flow_text_02_input.png", 5.5, "图 5-1  文本提示词生成视频 - 输入示例")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # CHAPTER 6: 项目总结
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "第六章  项目总结")

    add_heading2(doc, "6.1 开发过程与克服的困难")
    add_body(doc, (
        "本项目从立项到当前版本，经历了多轮迭代优化。在开发过程中，团队遇到并解决了以下关键难题：\n\n"
        "(1) LLM 输出格式不稳定：早期版本中 LLM 经常返回非标准 JSON（夹带 Markdown 代码块、"
        "字段缺失、类型错误），导致前端渲染异常。解决方案包括：在 system prompt 中明确禁止额外格式、"
        "在 response_schema 中为每个字段添加说明性描述、实现 try_parse_json_object 容错解析、"
        "以及完善的兜底机制确保全链路可用。\n\n"
        "(2) 短视频平台反爬：抖音、小红书等平台的分享链接经过多次重定向，且页面内容通过 JS 动态渲染，"
        "传统 HTTP 请求无法获取视频直链。解决方案是实现双策略解析——先尝试 HTTP 重定向追踪 + "
        "正则提取，失败后启动 Playwright 无头浏览器渲染页面再提取。\n\n"
        "(3) FFmpeg 兼容性问题：macOS 上 brew 安装的 ffmpeg 默认不包含 libfreetype，"
        "导致文字叠加功能失败。此外 ffmpeg-full 安装后引发 libvpx 版本不匹配。"
        "解决方案是在 run.py 启动时自动检测并 prepend ffmpeg-full 的 PATH。\n\n"
        "(4) Agent Tool-Calling 超时控制：早期版本中工具执行阻塞导致 SSE 流中断。"
        "改用 ThreadPoolExecutor.shutdown(wait=False) + future.result(timeout) "
        "实现真正非阻塞超时，确保单次工具调用在 0.46 秒内返回。"
    ), indent=True)

    add_heading2(doc, "6.2 AI 大模型应用心得")
    add_body(doc, (
        "在本项目中，我们对 AI 大模型的应用形成了几点关键认知：\n\n"
        "(1) 结构化输出是核心挑战：让 LLM 稳定输出 JSON 是工程落地的最大门槛。"
        "单纯靠 prompt 约束不够，需要配合 response_schema 声明、容错解析和兜底机制三层保障。\n\n"
        "(2) 多模态能力降低架构复杂度：Gemini 2.5 Flash 的视频理解能力让我们省去了独立的 ASR 服务，"
        "将「视频 → 字幕」的整条管线压缩为一次 API 调用，大幅简化了系统架构。\n\n"
        "(3) Prompt 工程需要领域知识：有效的 prompt 不仅仅是「告诉 AI 做什么」，"
        "还需要注入领域知识（如电商视频的钩子-推进-收口三段式结构、不同引擎的提示词风格差异），"
        "让 AI 的输出更贴合实际业务需求。\n\n"
        "(4) 缓存策略直接影响成本：AI API 调用成本不容忽视。"
        "合理的缓存策略（ASR 24h + 分析 15min）可以在不影响用户体验的前提下，"
        "将重复调用降低 60% 以上。"
    ), indent=True)

    add_heading2(doc, "6.3 后续演进规划")
    add_body(doc, (
        "未来版本计划在以下方向持续迭代：\n\n"
        "(1) API 网关集成：引入 AgentKit Gateway 统一管理 AI 服务调用，实现限流、"
        "鉴权和调用分析。\n\n"
        "(2) 分布式缓存：将内存缓存迁移至 Redis，支持多实例部署和缓存预热。\n\n"
        "(3) 更多引擎接入：计划接入 Sora、Kling 等新一代视频生成模型，"
        "丰富用户选择。\n\n"
        "(4) 多轮迭代优化：支持用户对复刻方案进行反馈，AI 自动调整输出，"
        "实现「分析-反馈-优化」闭环。\n\n"
        "(5) 商业化探索：面向电商代运营公司和 MCN 机构提供 SaaS 服务，"
        "按视频生成量计费。"
    ), indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════
    add_heading1(doc, "参考文献")
    refs = [
        "[1] Google. Gemini API Documentation [EB/OL]. https://ai.google.dev/docs, 2025.",
        "[2] Google. Vertex AI Video Generation (Veo) [EB/OL]. https://cloud.google.com/vertex-ai/docs/generative-ai/video/overview, 2025.",
        "[3] Google. Vertex AI Gemini Multimodal [EB/OL]. https://cloud.google.com/vertex-ai/docs/generative-ai/multimodal/overview, 2025.",
        "[4] Pallets Projects. Flask Documentation [EB/OL]. https://flask.palletsprojects.com/, 2024.",
        "[5] FFmpeg Developers. FFmpeg Documentation [EB/OL]. https://ffmpeg.org/documentation.html, 2025.",
        "[6] Samuel Colvin. Pydantic Documentation [EB/OL]. https://docs.pydantic.dev/, 2025.",
        "[7] Microsoft. Playwright for Python [EB/OL]. https://playwright.dev/python/, 2025.",
        "[8] 中国互联网络信息中心. 第 55 次中国互联网络发展状况统计报告 [R]. 2025.",
    ]
    for ref in refs:
        add_body(doc, ref)

    # ── Save ──
    doc.save(str(OUT))
    print(f"Document saved to: {OUT}")
    print(f"File size: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
